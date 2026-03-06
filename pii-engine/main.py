"""
Pii Sanitizer — PII Detection & Redaction Engine

FastAPI service that provides:
  POST /analyze  — detect PII entities in a file (PDF, DOCX, TXT, SQL)
  POST /sanitize — redact PII and return the sanitized file bytes
  GET  /health   — health check

Powered by Microsoft Presidio + spaCy.
"""

import io
import json
import logging
import re
import time
from collections import Counter
from typing import Optional

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response

from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer, RecognizerResult
from presidio_analyzer.nlp_engine import NlpEngineProvider
from presidio_analyzer.predefined_recognizers.country_specific.india import (
    InAadhaarRecognizer,
    InGstinRecognizer,
    InPanRecognizer,
    InPassportRecognizer,
    InVehicleRegistrationRecognizer,
    InVoterRecognizer,
)
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig

import pdfplumber
import fitz  # PyMuPDF — for in-place PDF redaction
from docx import Document

import csv as csv_mod

from parsers import extract_text, get_image_ocr_cache
from PIL import Image, ImageDraw

# ─── Logging ─────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("pii-engine")

# ─── Initialize engines ONCE at startup ──────────────────────────────────────

SPACY_MODEL = os.environ.get("SPACY_MODEL", "en_core_web_sm")
logger.info(f"Loading spaCy model ({SPACY_MODEL}) ...")

nlp_config = {
    "nlp_engine_name": "spacy",
    "models": [{"lang_code": "en", "model_name": SPACY_MODEL}],
}
nlp_engine = NlpEngineProvider(nlp_configuration=nlp_config).create_engine()

analyzer = AnalyzerEngine(nlp_engine=nlp_engine, supported_languages=["en"])

# ─── Custom recognizers ──────────────────────────────────────────────────────

# Indian phone number:  +91-XXXXXXXXXX, +91 XXXXXXXXXX, 0XX-XXXXXXXX, 10-digit mobile
indian_phone_recognizer = PatternRecognizer(
    supported_entity="PHONE_NUMBER",
    name="IndianPhoneRecognizer",
    patterns=[
        Pattern("IN_PHONE_INTL", r"(?:\+91[\-\s]?)?[6-9]\d{9}\b", 0.7),
        Pattern("IN_PHONE_LANDLINE", r"\b0\d{2,4}[\-\s]?\d{6,8}\b", 0.5),
    ],
    context=["phone", "mobile", "contact", "call", "tel", "number"],
    supported_language="en",
)

# Aadhaar — only match exactly 12 digits in 3 groups of 4 (not 4+ groups like credit cards)
aadhaar_recognizer = PatternRecognizer(
    supported_entity="IN_AADHAAR",
    name="EnhancedAadhaarRecognizer",
    patterns=[
        # Exactly 3 groups of 4 digits, NOT followed by another digit group
        Pattern("AADHAAR_FORMATTED", r"\b\d{4}[\s\-:]\d{4}[\s\-:]\d{4}\b(?![\s\-:]\d)", 0.6),
        Pattern("AADHAAR_PLAIN", r"\b\d{12}\b", 0.3),
    ],
    context=["aadhaar", "aadhar", "uid", "uidai", "unique identification"],
    supported_language="en",
)

# US SSN
us_ssn_recognizer = PatternRecognizer(
    supported_entity="US_SSN",
    name="UsSsnRecognizer",
    patterns=[
        Pattern("SSN_DASHES", r"\b\d{3}-\d{2}-\d{4}\b", 0.7),
        Pattern("SSN_SPACES", r"\b\d{3}\s\d{2}\s\d{4}\b", 0.5),
    ],
    context=["ssn", "social security", "social security number"],
    supported_language="en",
)

# Credit card — with Luhn validation
from presidio_analyzer import LocalRecognizer

class CreditCardPatternRecognizer(LocalRecognizer):
    """Detect credit card numbers in various formats with Luhn checksum."""

    PATTERNS = [
        Pattern("CC_DASHES", r"(?<![0-9])\d{4}-\d{4}-\d{4}-\d{4}(?![0-9])", 0.5),
        Pattern("CC_SPACES", r"(?<![0-9])\d{4}\s\d{4}\s\d{4}\s\d{4}(?![0-9])", 0.5),
        Pattern("CC_PLAIN", r"\b\d{13,19}\b", 0.1),
    ]

    def __init__(self):
        super().__init__(
            supported_entities=["CREDIT_CARD"],
            supported_language="en",
            name="CustomCreditCardRecognizer",
        )

    @staticmethod
    def _luhn_check(number: str) -> bool:
        digits = [int(d) for d in number if d.isdigit()]
        if len(digits) < 13 or len(digits) > 19:
            return False
        checksum = 0
        reverse_digits = digits[::-1]
        for i, d in enumerate(reverse_digits):
            if i % 2 == 1:
                d *= 2
                if d > 9:
                    d -= 9
            checksum += d
        return checksum % 10 == 0

    def load(self):
        pass

    def analyze(self, text, entities, nlp_artifacts):
        import re
        results = []
        for pattern in self.PATTERNS:
            for match in re.finditer(pattern.regex, text):
                raw = match.group()
                if self._luhn_check(raw):
                    score = 1.0  # Luhn-valid = high confidence
                else:
                    score = max(pattern.score, 0.4)  # format match = moderate confidence
                results.append(
                    RecognizerResult(
                        entity_type="CREDIT_CARD",
                        start=match.start(),
                        end=match.end(),
                        score=score,
                    )
                )
        return results

credit_card_recognizer = CreditCardPatternRecognizer()

# CVV / CVC — 3 or 4 digit security code (context-dependent)
cvv_recognizer = PatternRecognizer(
    supported_entity="CVV",
    name="CvvRecognizer",
    patterns=[
        Pattern("CVV_3", r"\b\d{3}\b", 0.3),
        Pattern("CVV_4", r"\b\d{4}\b", 0.2),
    ],
    context=["cvv", "cvc", "security code", "card verification", "cvv2", "cvc2"],
    supported_language="en",
)

# Lenient Indian PAN — matches any AAAAA9999A (relaxed 4th-char check)
lenient_pan_recognizer = PatternRecognizer(
    supported_entity="IN_PAN",
    name="LenientPanRecognizer",
    patterns=[
        Pattern("PAN_LENIENT", r"\b[A-Z]{5}\d{4}[A-Z]\b", 0.5),
    ],
    context=["pan", "permanent account", "income tax", "pan card", "pan number"],
    supported_language="en",
)

# Indian UPI ID — user@provider format
upi_recognizer = PatternRecognizer(
    supported_entity="UPI_ID",
    name="UpiRecognizer",
    patterns=[
        Pattern("UPI", r"\b[a-zA-Z0-9.\-_]+@(?:paytm|upi|okaxis|okhdfcbank|oksbi|ybl|ibl|apl|axl|icici|kotak|sbi|hdfcbank|axisbank|indus|freecharge|phonepe|gpay|amazonpay)\b", 0.8),
    ],
    context=["upi", "vpa", "payment"],
    supported_language="en",
)

# Indian IFSC code — 4-letter bank code + 0 + 6-char branch code
ifsc_recognizer = PatternRecognizer(
    supported_entity="IFSC",
    name="IfscRecognizer",
    patterns=[
        Pattern("IFSC", r"\b[A-Z]{4}0[A-Z0-9]{6}\b", 0.6),
    ],
    context=["ifsc", "bank", "branch", "neft", "rtgs", "imps"],
    supported_language="en",
)

# Register India-specific recognizers (built-in)
india_recognizers = [
    InAadhaarRecognizer(),
    InPanRecognizer(),
    InGstinRecognizer(),
    InPassportRecognizer(),
    InVehicleRegistrationRecognizer(),
    InVoterRecognizer(),
]
for rec in india_recognizers:
    analyzer.registry.add_recognizer(rec)

# Register custom recognizers
analyzer.registry.add_recognizer(indian_phone_recognizer)
analyzer.registry.add_recognizer(aadhaar_recognizer)
analyzer.registry.add_recognizer(us_ssn_recognizer)
analyzer.registry.add_recognizer(credit_card_recognizer)
analyzer.registry.add_recognizer(cvv_recognizer)
analyzer.registry.add_recognizer(lenient_pan_recognizer)
analyzer.registry.add_recognizer(upi_recognizer)
analyzer.registry.add_recognizer(ifsc_recognizer)

anonymizer = AnonymizerEngine()

logger.info("Engines ready. Supported entities: %s", analyzer.get_supported_entities())

# ─── FastAPI app ─────────────────────────────────────────────────────────────

app = FastAPI(
    title="Pii Sanitizer Engine",
    version="1.0.0",
    description="PII detection and redaction for PDF, DOCX, TXT, and SQL files.",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Minimum confidence threshold ────────────────────────────────────────────

SCORE_THRESHOLD = 0.35

# ─── Helper: deduplicate / merge overlapping results ─────────────────────────

# Higher priority types win when spans overlap (higher number = higher priority)
ENTITY_PRIORITY = {
    "CREDIT_CARD": 12,
    "IN_GSTIN": 11,
    "US_SSN": 10,
    "PHONE_NUMBER": 9,
    "EMAIL_ADDRESS": 9,
    "IP_ADDRESS": 9,
    "IN_PAN": 8,
    "IN_VEHICLE_REGISTRATION": 8,
    "IFSC": 8,
    "UPI_ID": 8,
    "IN_PASSPORT": 7,
    "IN_AADHAAR": 6,
    "IN_VOTER": 6,
    "PERSON": 5,
    "LOCATION": 4,
    "CVV": 3,
    "ORGANIZATION": 3,
    "DATE_TIME": 2,
    "NRP": 1,
}


def _rank(r: RecognizerResult) -> tuple:
    """Rank a result: prefer type priority, then score, then longer span."""
    span_len = r.end - r.start
    priority = ENTITY_PRIORITY.get(r.entity_type, 0)
    return (priority, r.score, span_len)


def deduplicate(results: list[RecognizerResult]) -> list[RecognizerResult]:
    """Remove overlapping detections — keep longer, higher-priority match."""
    if not results:
        return []
    sorted_results = sorted(results, key=lambda r: r.start)
    merged: list[RecognizerResult] = [sorted_results[0]]
    for current in sorted_results[1:]:
        prev = merged[-1]
        if current.start < prev.end:
            # Overlap — keep the one that ranks higher
            if _rank(current) > _rank(prev):
                merged[-1] = current
        else:
            merged.append(current)
    return merged


# ─── False-positive filter ────────────────────────────────────────────────

# Entity types that are almost always false positives (document titles, labels)
_SKIP_ENTITY_TYPES = {"ORGANIZATION", "NRP"}

# Common form-field labels — if the entire entity text matches one of these
# (case-insensitive, stripped), it is a label, not actual PII data.
_LABEL_WORDS = {
    "name", "email", "phone", "address", "aadhaar", "pan", "ifsc", "upi",
    "cvv", "credit card", "bank account", "ip address", "ip_address",
    "field", "example", "example pii", "pii", "dob", "date of birth",
    "mobile", "contact", "passport", "voter id", "gstin",
    "vehicle registration", "upi_id", "ifsc_code",
    "first_name", "last_name", "full_name", "first name", "last name",
    "full name", "card_number", "card number", "expiration_date",
    "phone_number", "phone number", "email_address", "email address",
    "credit_card", "ip_address", "address_text",
}

# Minimum text length per entity type to avoid matching short labels
_MIN_ENTITY_LENGTHS = {
    "IN_PAN": 10,
    "IN_AADHAAR": 12,
    "CREDIT_CARD": 13,
    "PHONE_NUMBER": 7,
    "EMAIL_ADDRESS": 5,
    "IP_ADDRESS": 7,
    "IN_GSTIN": 15,
    "IN_PASSPORT": 8,
    "US_SSN": 9,
}


def _should_skip_entity(ent_type: str, ent_text: str) -> bool:
    """Return True if this entity is a false positive."""
    if ent_type in _SKIP_ENTITY_TYPES:
        return True
    stripped = ent_text.strip()
    if stripped.lower() in _LABEL_WORDS:
        return True
    min_len = _MIN_ENTITY_LENGTHS.get(ent_type)
    if min_len and len(stripped) < min_len:
        return True
    return False


def _trim_person_entity(ent_text: str, start: int, end: int):
    """For PERSON entities, strip trailing/leading label words that spaCy
    incorrectly absorbed (e.g. 'Rahul Sharma\nEmail' → 'Rahul Sharma').
    Returns (text, start, end) or None to skip entirely."""
    tokens = list(re.finditer(r'\S+', ent_text))
    if not tokens:
        return None
    good = [t for t in tokens if t.group().lower() not in _LABEL_WORDS]
    if not good:
        return None
    first, last = good[0], good[-1]
    new_text = ent_text[first.start():last.end()]
    return (new_text, start + first.start(), start + last.end())


def _clean_results(results: list[RecognizerResult], text: str) -> list[RecognizerResult]:
    """Filter false positives and trim label words from PERSON entities."""
    cleaned: list[RecognizerResult] = []
    for r in results:
        ent_text = text[r.start:r.end]
        if _should_skip_entity(r.entity_type, ent_text):
            continue
        if r.entity_type == "PERSON":
            trimmed = _trim_person_entity(ent_text, r.start, r.end)
            if trimmed is None:
                continue
            new_text, new_start, new_end = trimmed
            if (new_start, new_end) != (r.start, r.end):
                r = RecognizerResult(
                    entity_type=r.entity_type,
                    start=new_start,
                    end=new_end,
                    score=r.score,
                )
        cleaned.append(r)
    return cleaned


def _clean_entity_dicts(entity_list: list[dict]) -> list[dict]:
    """Filter false positives and trim label words from PERSON entity dicts."""
    cleaned: list[dict] = []
    for ent in entity_list:
        ent_type = ent.get("type", "")
        ent_text = ent.get("text", "")
        if _should_skip_entity(ent_type, ent_text):
            continue
        if ent_type == "PERSON":
            trimmed = _trim_person_entity(ent_text, ent["start"], ent["end"])
            if trimmed is None:
                continue
            new_text, new_start, new_end = trimmed
            if new_text != ent_text:
                ent = dict(ent)
                ent["text"] = new_text
                ent["start"] = new_start
                ent["end"] = new_end
        cleaned.append(ent)
    return cleaned


# ─── Masking helpers ─────────────────────────────────────────────────────────

# Entity types whose values should stay VISIBLE (not masked).
# They are still detected and reported, but the sanitized output keeps them.
_KEEP_VISIBLE_TYPES = {"PERSON", "EMAIL_ADDRESS", "LOCATION"}


def _mask_entity(entity_type: str, original_text: str) -> str | None:
    """Return a masked version of the PII text.  Returns None for types
    that should stay visible (names, emails, locations)."""

    if entity_type in _KEEP_VISIBLE_TYPES:
        return None  # keep original — do not redact

    if entity_type == "CVV":
        # "489" → "***"
        return re.sub(r'\d', '*', original_text)

    if entity_type == "PHONE_NUMBER":
        # Mask last 2 digits only: "(217) 555-0123" → "(217) 555-01XX"
        digits = list(re.finditer(r'\d', original_text))
        if len(digits) > 2:
            result = list(original_text)
            for d in digits[-2:]:
                result[d.start()] = 'X'
            return ''.join(result)
        return re.sub(r'\d', 'X', original_text)

    if entity_type == "US_SSN":
        # "321-45-7890" → "XXX-XX-7890" (show last 4)
        digits = list(re.finditer(r'\d', original_text))
        result = list(original_text)
        for d in digits[:-4]:
            result[d.start()] = 'X'
        return ''.join(result)

    if entity_type == "IN_AADHAAR":
        # "1234 5678 9012" → "XXXX XXXX 9012" (show last 4)
        digits = list(re.finditer(r'\d', original_text))
        result = list(original_text)
        for d in digits[:-4]:
            result[d.start()] = 'X'
        return ''.join(result)

    if entity_type == "IN_PAN":
        # "ABCDE1234F" → "XXXXXXXXXX"
        return re.sub(r'[a-zA-Z0-9]', 'X', original_text)

    if entity_type == "CREDIT_CARD":
        # "4111-1111-1111-1111" → "****-****-****-1111" (show last 4)
        digits = list(re.finditer(r'\d', original_text))
        result = list(original_text)
        for d in digits[:-4]:
            result[d.start()] = '*'
        # Replace separators between masked groups with spaces for clean look
        return ''.join(result)

    if entity_type == "US_BANK_NUMBER":
        # "123456789012" → "XXXXXXXX9012" (show last 4)
        digits = list(re.finditer(r'\d', original_text))
        result = list(original_text)
        for d in digits[:-4]:
            result[d.start()] = 'X'
        return ''.join(result)

    if entity_type == "IP_ADDRESS":
        # "192.168.1.105" → "XXX.XXX.X.XXX"
        return re.sub(r'\d', 'X', original_text)

    if entity_type == "DATE_TIME":
        # "May 12, 1985" → "May 12, 19XX" (mask last 2 digits only)
        digits = list(re.finditer(r'\d', original_text))
        if len(digits) > 2:
            result = list(original_text)
            for d in digits[-2:]:
                result[d.start()] = 'X'
            return ''.join(result)
        return re.sub(r'\d', 'X', original_text)

    if entity_type == "UPI_ID":
        # "rahul@paytm" → "XXXXX@XXXXXX"
        return re.sub(r'[a-zA-Z0-9]', 'X', original_text)

    if entity_type == "IFSC":
        # "HDFC0001234" → "XXXXXXXXXXX"
        return re.sub(r'[a-zA-Z0-9]', 'X', original_text)

    # Default: mask all alphanumeric
    return re.sub(r'[a-zA-Z0-9]', 'X', original_text)


def _mask_text_manual(text: str, entity_list: list[dict]) -> str:
    """Replace PII in text with masked versions (offset-based, end-to-start).
    Skips entities whose type is in _KEEP_VISIBLE_TYPES."""
    sorted_ents = sorted(entity_list, key=lambda e: e["start"], reverse=True)
    result = text
    for ent in sorted_ents:
        masked = _mask_entity(ent["type"], ent["text"])
        if masked is None:
            continue  # keep visible
        result = result[:ent["start"]] + masked + result[ent["end"]:]
    return result


# ─── Structured-data CVV detection ───────────────────────────────────────────
# In CSV/SQL/JSON the column header "cvv" is too far from the cell values
# for Presidio's context window.  These helpers find CVV values by column/key.

_CVV_COLUMN_NAMES = {"cvv", "cvc", "cvv2", "cvc2", "security_code", "card_verification"}


def _csv_cell_spans(line: str) -> list[tuple[int, int]]:
    """Return (start, end) byte offsets for each cell in a CSV line."""
    spans: list[tuple[int, int]] = []
    start = 0
    in_quotes = False
    for i, ch in enumerate(line):
        if ch == '"':
            in_quotes = not in_quotes
        elif ch == ',' and not in_quotes:
            spans.append((start, i))
            start = i + 1
    spans.append((start, len(line)))
    return spans


def _detect_csv_column_entities(text: str) -> list[RecognizerResult]:
    """Find CVV values in CSV by column-header mapping."""
    lines = text.split('\n')
    if len(lines) < 2:
        return []
    try:
        header = next(csv_mod.reader(io.StringIO(lines[0])))
    except Exception:
        return []
    cvv_cols: set[int] = set()
    for i, h in enumerate(header):
        if h.strip().lower() in _CVV_COLUMN_NAMES:
            cvv_cols.add(i)
    if not cvv_cols:
        return []
    results: list[RecognizerResult] = []
    offset = len(lines[0]) + 1
    for line_idx in range(1, len(lines)):
        line = lines[line_idx]
        if not line.strip():
            offset += len(line) + 1
            continue
        cell_spans = _csv_cell_spans(line)
        for col_idx in cvv_cols:
            if col_idx >= len(cell_spans):
                continue
            s, e = cell_spans[col_idx]
            cell_text = line[s:e]
            m = re.search(r'\d{3,4}', cell_text)
            if m:
                results.append(RecognizerResult(
                    entity_type="CVV", start=offset + s + m.start(),
                    end=offset + s + m.end(), score=0.85))
        offset += len(line) + 1
    return results


def _sql_value_spans(values_text: str) -> list[tuple[int, int]]:
    """Return (start, end) offsets for each value in a SQL VALUES clause."""
    spans: list[tuple[int, int]] = []
    start = 0
    i = 0
    in_quotes = False
    quote_char = None
    while i < len(values_text):
        ch = values_text[i]
        if not in_quotes:
            if ch in ("'", '"'):
                in_quotes = True
                quote_char = ch
            elif ch == ',':
                spans.append((start, i))
                start = i + 1
        else:
            if ch == quote_char:
                if i + 1 < len(values_text) and values_text[i + 1] == quote_char:
                    i += 2
                    continue
                in_quotes = False
        i += 1
    spans.append((start, len(values_text)))
    return spans


def _detect_sql_column_entities(text: str) -> list[RecognizerResult]:
    """Find CVV values in SQL INSERT statements by column order.
    Supports both:
      - CREATE TABLE ... (col1 TYPE, col2 TYPE) with VALUES (...);
      - INSERT INTO table (col1, col2) VALUES (...), (...), ...;
    """
    cvv_indices: set[int] = set()
    columns: list[str] = []

    # Strategy 1: parse columns from INSERT INTO ... (columns) VALUES
    for ins_match in re.finditer(
        r'INSERT\s+INTO\s+\w+\s*\(([^)]+)\)\s*VALUES',
        text, re.IGNORECASE,
    ):
        cols = [c.strip().lower() for c in ins_match.group(1).split(',')]
        for i, c in enumerate(cols):
            if c in _CVV_COLUMN_NAMES:
                cvv_indices.add(i)
        if cvv_indices:
            columns = cols
            break

    # Strategy 2 fallback: parse from CREATE TABLE
    if not cvv_indices:
        create_match = re.search(
            r'CREATE\s+TABLE\s+\w+\s*\((.*?)\)\s*;', text,
            re.IGNORECASE | re.DOTALL)
        if create_match:
            columns = [col_def.strip().split()[0].strip().lower()
                       for col_def in re.split(r',(?![^(]*\))', create_match.group(1))]
            cvv_indices = {i for i, c in enumerate(columns) if c in _CVV_COLUMN_NAMES}

    if not cvv_indices:
        return []

    results: list[RecognizerResult] = []
    # Match individual value tuples: (val1, val2, ...) anywhere after VALUES
    for match in re.finditer(r'\(([^()]+)\)', text):
        # Only consider tuples that appear after a VALUES keyword
        before = text[max(0, match.start() - 200):match.start()]
        if not re.search(r'VALUES\s*$', before, re.IGNORECASE) and \
           not re.search(r'\)\s*,\s*$', before):
            continue
        v_text = match.group(1)
        v_start = match.start(1)
        val_spans = _sql_value_spans(v_text)
        for col_idx in cvv_indices:
            if col_idx >= len(val_spans):
                continue
            vs, ve = val_spans[col_idx]
            cell = v_text[vs:ve].strip().strip("'\"")
            if not cell or cell.upper() == 'NULL':
                continue
            dm = re.search(r'\d{3,4}', v_text[vs:ve])
            if dm:
                results.append(RecognizerResult(
                    entity_type="CVV", start=v_start + vs + dm.start(),
                    end=v_start + vs + dm.end(), score=0.85))
    return results


def _detect_json_key_entities(text: str) -> list[RecognizerResult]:
    """Find CVV values in JSON by key names."""
    results: list[RecognizerResult] = []
    for key_name in _CVV_COLUMN_NAMES:
        pattern = rf'["\']({re.escape(key_name)})["\']\s*:\s*["\']?(\d{{3,4}})["\']?'
        for m in re.finditer(pattern, text, re.IGNORECASE):
            results.append(RecognizerResult(
                entity_type="CVV", start=m.start(2),
                end=m.end(2), score=0.85))
    return results


def _detect_structured_cvv(
    text: str, content_type: str, filename: str,
) -> list[RecognizerResult]:
    """Detect CVV values in structured data where Presidio's context window
    cannot reach from the column/key header to the cell value."""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if content_type in ('text/csv', 'application/csv') or ext == 'csv':
        return _detect_csv_column_entities(text)
    if content_type == 'application/json' or ext == 'json':
        return _detect_json_key_entities(text)
    if content_type in ('text/sql', 'application/sql', 'application/x-sql') or ext == 'sql':
        return _detect_sql_column_entities(text)
    return []


def _filter_cvv_false_positives(
    results: list[RecognizerResult], text: str,
) -> list[RecognizerResult]:
    """Remove CVV detections that are actually part of a date or IP pattern.
    Presidio's context window can bleed from a nearby 'cvv' key to an
    adjacent value, causing digits to be mis-classified as CVV."""
    filtered: list[RecognizerResult] = []
    for r in results:
        if r.entity_type == "CVV":
            # Date patterns DD/MM/YYYY or YYYY-MM-DD
            ctx_before = text[max(0, r.start - 6):r.end]
            ctx_after = text[r.start:min(len(text), r.end + 6)]
            if re.search(r'\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{3,4}$', ctx_before):
                continue
            if re.search(r'^\d{4}[/\-\.]\d{1,2}[/\-\.]', ctx_after):
                continue
            # IP address: only reject if the CVV span overlaps an IP match
            ctx_start = max(0, r.start - 16)
            ctx_ip = text[ctx_start:min(len(text), r.end + 16)]
            is_in_ip = False
            for m in re.finditer(r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}', ctx_ip):
                ip_abs_start = ctx_start + m.start()
                ip_abs_end = ctx_start + m.end()
                if r.start >= ip_abs_start and r.end <= ip_abs_end:
                    is_in_ip = True
                    break
            if is_in_ip:
                continue
        filtered.append(r)
    return filtered


# ═════════════════════════════════════════════════════════════════════════════
#  POST /analyze
# ═════════════════════════════════════════════════════════════════════════════

@app.post("/analyze")
async def analyze(file: UploadFile = File(...)):
    """
    Detect PII entities in the uploaded file.

    Returns:
        {
          "entities": [ { "type", "text", "start", "end", "score" }, ... ],
          "stats": { "total": N, "by_type": { "EMAIL_ADDRESS": 2, ... } }
        }
    """
    start_time = time.time()

    # Read file
    data = await file.read()
    filename = file.filename or "unknown"
    content_type = file.content_type or "application/octet-stream"

    logger.info("Analyze request — file=%s, type=%s, size=%d bytes", filename, content_type, len(data))

    # Extract text
    try:
        text = extract_text(data, content_type, filename)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    logger.info("Extracted %d characters of text", len(text))

    # Run Presidio analyzer
    raw_results = analyzer.analyze(
        text=text,
        language="en",
        score_threshold=SCORE_THRESHOLD,
    )
    # Add column/key-aware CVV detection for structured formats
    raw_results.extend(_detect_structured_cvv(text, content_type, filename))
    raw_results = _filter_cvv_false_positives(raw_results, text)

    results = deduplicate(raw_results)

    # Filter out false positives and trim label words from PERSON entities
    results = _clean_results(results, text)

    # Build response
    entities = []
    for r in results:
        entities.append({
            "type": r.entity_type,
            "text": text[r.start : r.end],
            "start": r.start,
            "end": r.end,
            "score": round(r.score, 4),
        })

    by_type: dict[str, int] = dict(Counter(e["type"] for e in entities))

    elapsed = round(time.time() - start_time, 3)
    logger.info(
        "Analyze complete — %d entities found in %.3fs (types: %s)",
        len(entities), elapsed, by_type,
    )

    return {
        "entities": entities,
        "stats": {
            "total": len(entities),
            "by_type": by_type,
        },
    }


# ═════════════════════════════════════════════════════════════════════════════
#  POST /sanitize
# ═════════════════════════════════════════════════════════════════════════════

@app.post("/sanitize")
async def sanitize(
    file: UploadFile = File(...),
    entities: str = Form(...),
):
    """
    Redact PII from the uploaded file and return the sanitized version.

    Expects:
      - file: the original file
      - entities: JSON string — the entities array from /analyze

    Returns: file bytes (same format as input) with PII replaced by [ENTITY_TYPE]
    """
    start_time = time.time()

    data = await file.read()
    filename = file.filename or "unknown"
    content_type = file.content_type or "application/octet-stream"

    logger.info("Sanitize request — file=%s, type=%s", filename, content_type)

    # Parse entities
    try:
        entity_list = json.loads(entities)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid entities JSON")

    # Filter out false positives and trim label words from PERSON entities
    entity_list = _clean_entity_dicts(entity_list)

    ext = ""
    if "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()

    output_bytes: bytes
    output_mime: str

    if content_type == "application/pdf" or ext == "pdf":
        output_bytes = _sanitize_pdf_inplace(data, entity_list)
        output_mime = "application/pdf"

    elif (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or ext == "docx"
    ):
        output_bytes = _sanitize_docx_inplace(data, entity_list)
        output_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif content_type in ("image/png", "image/jpeg") or ext in ("png", "jpg", "jpeg"):
        output_bytes = _sanitize_image(data, entity_list)
        output_mime = "image/png"

    else:
        # Text-based formats (SQL, TXT, CSV, etc.) — mask with X characters
        try:
            text = extract_text(data, content_type, filename)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))

        redacted_text = _mask_text_manual(text, entity_list)

        if ext == "sql" or content_type in ("text/sql", "application/sql", "application/x-sql"):
            output_bytes = redacted_text.encode("utf-8")
            output_mime = "text/sql"
        else:
            output_bytes = redacted_text.encode("utf-8")
            output_mime = "text/plain"

    elapsed = round(time.time() - start_time, 3)
    logger.info(
        "Sanitize complete — %d entities redacted in %.3fs, output %d bytes",
        len(entity_list), elapsed, len(output_bytes),
    )

    # Fix filename extension for image outputs
    out_filename = filename
    if output_mime == "image/png" and not filename.lower().endswith(".png"):
        out_filename = filename.rsplit(".", 1)[0] + ".png" if "." in filename else filename + ".png"

    return Response(
        content=output_bytes,
        media_type=output_mime,
        headers={
            "Content-Disposition": f'attachment; filename="sanitized_{out_filename}"',
        },
    )


# ═════════════════════════════════════════════════════════════════════════════
#  POST /analyze-and-sanitize  (convenience: one call does both)
# ═════════════════════════════════════════════════════════════════════════════

@app.post("/analyze-and-sanitize")
async def analyze_and_sanitize(file: UploadFile = File(...)):
    """
    Combined endpoint: analyze + sanitize in one round trip.
    Returns JSON with stats AND the sanitized text (base64 for binary formats).
    """
    import base64

    data = await file.read()
    filename = file.filename or "unknown"
    content_type = file.content_type or "application/octet-stream"

    try:
        text = extract_text(data, content_type, filename)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    raw_results = analyzer.analyze(text=text, language="en", score_threshold=SCORE_THRESHOLD)
    # Add column/key-aware CVV detection for structured formats
    raw_results.extend(_detect_structured_cvv(text, content_type, filename))
    raw_results = _filter_cvv_false_positives(raw_results, text)
    results = deduplicate(raw_results)

    # Filter out false positives and trim label words from PERSON entities
    results = _clean_results(results, text)

    entities = []
    for r in results:
        entities.append({
            "type": r.entity_type,
            "text": text[r.start : r.end],
            "start": r.start,
            "end": r.end,
            "score": round(r.score, 4),
        })

    sanitized_text = _mask_text_manual(text, entities)

    by_type = dict(Counter(e["type"] for e in entities))

    return {
        "entities": entities,
        "stats": {"total": len(entities), "by_type": by_type},
        "sanitized_text": sanitized_text,
    }


# ═════════════════════════════════════════════════════════════════════════════
#  GET /health
# ═════════════════════════════════════════════════════════════════════════════

@app.get("/health")
async def health():
    return {"status": "ok", "entities_supported": analyzer.get_supported_entities()}


# ═════════════════════════════════════════════════════════════════════════════
#  In-place DOCX sanitisation (preserves formatting)
# ═════════════════════════════════════════════════════════════════════════════

def _replace_text_preserving_format(paragraph, old_text: str, new_text: str):
    """
    Replace every occurrence of *old_text* with *new_text* across the
    paragraph's runs, keeping all run-level formatting (bold, italic,
    colour, font, size …).  The replacement inherits the style of the
    first run it touches.
    """
    while True:
        runs = paragraph.runs
        if not runs:
            break

        # Build a character → (run_index, char_index) map
        combined = ""
        char_map: list[tuple[int, int]] = []
        for i, run in enumerate(runs):
            for j in range(len(run.text)):
                char_map.append((i, j))
            combined += run.text

        idx = combined.find(old_text)
        if idx == -1:
            break

        end_idx = idx + len(old_text)
        if end_idx > len(char_map):
            break

        first_run_idx, first_char = char_map[idx]
        last_run_idx,  last_char  = char_map[end_idx - 1]

        if first_run_idx == last_run_idx:
            r = runs[first_run_idx]
            r.text = r.text[:first_char] + new_text + r.text[last_char + 1:]
        else:
            # First run: keep text before entity + insert replacement
            runs[first_run_idx].text = (
                runs[first_run_idx].text[:first_char] + new_text
            )
            # Middle runs: clear
            for mid in range(first_run_idx + 1, last_run_idx):
                runs[mid].text = ""
            # Last run: keep text after entity
            runs[last_run_idx].text = runs[last_run_idx].text[last_char + 1:]


def _sanitize_docx_inplace(data: bytes, entity_list: list) -> bytes:
    """
    Open the original DOCX, walk every paragraph/table-cell/header/footer,
    and replace PII text directly inside the existing XML runs so that
    bold, italic, colour, font, size and all other formatting is kept.

    Uses a two-phase approach (entity→placeholder→tag) so that replacement
    tags like [US_SSN] are never corrupted by shorter entity matches.
    """
    doc = Document(io.BytesIO(data))

    # Phase 1 map: entity text → unique placeholder (PUA Unicode)
    # Phase 2 map: placeholder → masked text
    # Skip keep-visible types (PERSON, EMAIL, LOCATION)
    replacement_map: dict[str, str] = {}
    placeholder_to_tag: dict[str, str] = {}
    idx = 0
    for ent in entity_list:
        txt = ent.get("text", "")
        masked = _mask_entity(ent.get("type", ""), txt)
        if txt and masked is not None and txt not in replacement_map:
            placeholder = f"\uE000{idx}\uE001"
            replacement_map[txt] = placeholder
            placeholder_to_tag[placeholder] = masked
            idx += 1

    sorted_phase1 = sorted(
        replacement_map.items(), key=lambda x: len(x[0]), reverse=True
    )
    sorted_phase2 = sorted(
        placeholder_to_tag.items(), key=lambda x: len(x[0]), reverse=True
    )

    def _process_paragraphs(paragraphs):
        for para in paragraphs:
            # Phase 1: entity text → placeholder
            for original, placeholder in sorted_phase1:
                if original in para.text:
                    _replace_text_preserving_format(para, original, placeholder)
            # Phase 2: placeholder → [TYPE] tag
            for placeholder, tag in sorted_phase2:
                if placeholder in para.text:
                    _replace_text_preserving_format(para, placeholder, tag)

    def _process_table(table):
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs)
                for nested in cell.tables:
                    _process_table(nested)

    # Body paragraphs
    _process_paragraphs(doc.paragraphs)

    # Body tables
    for table in doc.tables:
        _process_table(table)

    # Headers & footers (all section variants)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf is None:
                continue
            _process_paragraphs(hf.paragraphs)
            for table in hf.tables:
                _process_table(table)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════════
#  In-place PDF sanitisation (preserves layout / fonts / images)
# ═════════════════════════════════════════════════════════════════════════════

def _sanitize_pdf_inplace(data: bytes, entity_list: list) -> bytes:
    """
    Open the original PDF with PyMuPDF, search for each PII text on every
    page, add a redaction annotation with the replacement tag, then apply
    all redactions.  Layout, images, headers, and colours are preserved.
    """
    pdf = fitz.open(stream=data, filetype="pdf")

    # Unique replacement map (longer first) — skip keep-visible types
    replacement_map: dict[str, str] = {}
    for ent in entity_list:
        txt = ent.get("text", "")
        masked = _mask_entity(ent.get("type", ""), txt)
        if txt and masked is not None and txt not in replacement_map:
            replacement_map[txt] = masked

    sorted_items = sorted(
        replacement_map.items(), key=lambda x: len(x[0]), reverse=True
    )

    for page in pdf:
        for original, replacement in sorted_items:
            hits = page.search_for(original)
            for rect in hits:
                page.add_redact_annot(
                    rect,
                    text=replacement,
                    fontsize=0,           # auto-fit to rect
                    text_color=(0, 0, 0), # black text
                    fill=(1, 1, 1),       # white background
                )
        page.apply_redactions()

    buf = io.BytesIO()
    pdf.save(buf, garbage=4, deflate=True)
    pdf.close()
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════════
#  In-place Image sanitisation (black bars over PII words)
# ═════════════════════════════════════════════════════════════════════════════

# Entity types that are NOT personal PII on identity documents
# (e.g., "Income Tax Department", "Government of India" = card headers)
_IMAGE_SKIP_TYPES = {"ORGANIZATION", "NRP"}

# Regex patterns for PII that Presidio might miss on noisy OCR text
_IMAGE_PII_PATTERNS = [
    r'[A-Z]{5}\d{4}[A-Z]',                              # Indian PAN
    r'\d{4}\s\d{4}\s\d{4}(?!\s\d)',                     # Aadhaar (12 digits)
    r'\d{4}\s\d{4}\s\d{4}\s\d{4}',                      # VID (16 digits)
    r'(?:\+91[\s\-]?)?[6-9]\d{9}\b',                    # Indian phone
    r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}',  # Email
]


def _sanitize_image(data: bytes, entity_list: list) -> bytes:
    """
    Open the image, use cached OCR bounding boxes with character offsets
    to map entity spans to pixel regions, plus regex fallback for PII
    patterns Presidio may miss.  Draw black rectangles over PII words.
    Always outputs PNG.
    """
    img = Image.open(io.BytesIO(data)).convert("RGB")

    # Upscale — must match parse_image() formula exactly
    w, h = img.size
    scale = 1
    if w < 2000:
        scale = max(2, (2000 // w) + 1)
        img = img.resize((w * scale, h * scale), Image.LANCZOS)

    # Get cached OCR word data (with char_start / char_end offsets)
    ocr_words = get_image_ocr_cache(data)
    if not ocr_words:
        logger.warning("No cached OCR data found for image — re-running OCR")
        from parsers import parse_image as _parse_img
        _parse_img(data)
        ocr_words = get_image_ocr_cache(data)

    # ── 1. Offset-based matching from Presidio entities ──────────────────
    #    Skip ORGANIZATION / NRP which are card headers, not personal PII
    words_to_redact: set[int] = set()
    for ent in entity_list:
        if ent.get("type") in _IMAGE_SKIP_TYPES:
            continue
        ent_start = ent["start"]
        ent_end = ent["end"]
        for idx, w in enumerate(ocr_words):
            w_start = w.get("char_start", -1)
            w_end = w.get("char_end", -1)
            if w_start < ent_end and w_end > ent_start:
                words_to_redact.add(idx)

    # ── 2. Regex fallback — catch PAN / Aadhaar / phone / email that ─────
    #    Presidio may have missed due to noisy OCR text
    ocr_text = " ".join(w["text"] for w in ocr_words)
    for pattern in _IMAGE_PII_PATTERNS:
        for m in re.finditer(pattern, ocr_text):
            ms, me = m.start(), m.end()
            for idx, w in enumerate(ocr_words):
                ws = w.get("char_start", -1)
                we = w.get("char_end", -1)
                if ws < me and we > ms:
                    words_to_redact.add(idx)

    # ── 3. Draw black rectangles ─────────────────────────────────────────
    draw = ImageDraw.Draw(img)
    img_w, img_h = img.size
    padding = 4

    for idx in words_to_redact:
        word_info = ocr_words[idx]
        left = max(0, word_info["left"] - padding)
        top = max(0, word_info["top"] - padding)
        right = min(img_w, word_info["left"] + word_info["width"] + padding)
        bottom = min(img_h, word_info["top"] + word_info["height"] + padding)
        if right > left and bottom > top:
            draw.rectangle([left, top, right, bottom], fill="black")

    logger.info("Image redaction: %d words blacked out of %d OCR words", len(words_to_redact), len(ocr_words))

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════════
#  Run (for local testing: python main.py)
# ═════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import os
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
