"""
File parsers — extract plain text from PDF, DOCX, TXT, SQL, and image files.

Each parser returns a single string of the full text content.
The /analyze endpoint runs PII detection on this string.
The /sanitize endpoint uses character offsets into this same string.
"""

import io
import pdfplumber
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract

# ─── OCR cache: stores bounding-box data between analyze and sanitize calls ──
# key = hash of image bytes, value = list of {text, left, top, width, height, conf, char_start, char_end}
_image_ocr_cache: dict[int, list[dict]] = {}


def parse_pdf(data: bytes) -> str:
    """Extract all text from a PDF file."""
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def parse_docx(data: bytes) -> str:
    """Extract all text from a DOCX file (paragraphs + tables)."""
    doc = Document(io.BytesIO(data))

    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def parse_txt(data: bytes) -> str:
    """Decode plain text (UTF-8 with fallback to latin-1)."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1")


def parse_sql(data: bytes) -> str:
    """
    Parse SQL files — treat as plain text.
    SQL may contain PII in INSERT/UPDATE values, comments, etc.
    """
    return parse_txt(data)


def parse_image(data: bytes) -> str:
    """
    Extract text from an image using Tesseract OCR.
    Preprocesses for better accuracy on ID cards (grayscale, contrast, sharpen).
    Caches word-level bounding boxes WITH character offsets
    for precise entity-to-pixel mapping during sanitization.
    """
    img = Image.open(io.BytesIO(data)).convert("RGB")

    # Upscale small images aggressively for better OCR on card photos
    w, h = img.size
    scale = 1
    if w < 2000:
        scale = max(2, (2000 // w) + 1)
        img = img.resize((w * scale, h * scale), Image.LANCZOS)

    # Preprocess for OCR: grayscale → contrast enhancement → sharpen
    gray = img.convert("L")
    gray = ImageEnhance.Contrast(gray).enhance(1.5)
    gray = gray.filter(ImageFilter.SHARPEN)

    # Run OCR with bounding box data
    ocr_data = pytesseract.image_to_data(
        gray, config="--oem 3 --psm 4", output_type=pytesseract.Output.DICT
    )

    # Cache the bounding box data keyed by image bytes hash
    # Track character offset of each word in the joined text
    cache_key = hash(data)
    words = []
    current_offset = 0
    for i in range(len(ocr_data["text"])):
        word = ocr_data["text"][i].strip()
        conf = int(ocr_data["conf"][i])
        if word and conf > 0:
            words.append({
                "text": word,
                "left": ocr_data["left"][i],
                "top": ocr_data["top"][i],
                "width": ocr_data["width"][i],
                "height": ocr_data["height"][i],
                "conf": conf,
                "char_start": current_offset,
                "char_end": current_offset + len(word),
            })
            current_offset += len(word) + 1  # +1 for the space separator
    _image_ocr_cache[cache_key] = words

    # Return plain text for PII analysis (space-joined, matching offsets)
    text_parts = [w["text"] for w in words]
    return " ".join(text_parts)


def get_image_ocr_cache(data: bytes) -> list[dict]:
    """Retrieve cached OCR bounding-box data for an image."""
    return _image_ocr_cache.get(hash(data), [])


# ─── Dispatcher ──────────────────────────────────────────────────────────────

MIME_PARSERS = {
    "application/pdf": parse_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": parse_docx,
    "text/plain": parse_txt,
    "text/csv": parse_txt,
    "application/json": parse_txt,
    "text/xml": parse_txt,
    "application/xml": parse_txt,
    "text/sql": parse_sql,
    "application/sql": parse_sql,
    "application/x-sql": parse_sql,
    "image/png": parse_image,
    "image/jpeg": parse_image,
}

# Fallback: detect by file extension
EXT_PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".txt": parse_txt,
    ".sql": parse_sql,
    ".csv": parse_txt,
    ".log": parse_txt,
    ".json": parse_txt,
    ".xml": parse_txt,
    ".png": parse_image,
    ".jpg": parse_image,
    ".jpeg": parse_image,
}


def extract_text(data: bytes, content_type: str, filename: str) -> str:
    """
    Extract text from file bytes using MIME type or extension fallback.
    Raises ValueError for unsupported file types.
    """
    # Try by MIME type first
    parser = MIME_PARSERS.get(content_type)

    # Fallback: try by extension
    if parser is None:
        ext = ""
        if "." in filename:
            ext = "." + filename.rsplit(".", 1)[-1].lower()
        parser = EXT_PARSERS.get(ext)

    if parser is None:
        raise ValueError(
            f"Unsupported file type: {content_type} ({filename}). "
            f"Supported: PDF, DOCX, TXT, SQL"
        )

    text = parser(data)
    if not text or not text.strip():
        raise ValueError("Could not extract any text from the file.")

    return text
